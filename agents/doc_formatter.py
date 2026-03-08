"""Doc Formatter Agent

Generates a professional .docx from the final letter using python-docx.
Pure code — no LLM calls needed.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


class DocFormatter:
    """Generates professional .docx documents."""

    name = "Doc Formatter"

    def run(
        self,
        letter_text: str,
        client_name: str,
        output_path: str,
        report_type: str = "monthly_letter",
    ) -> dict:
        """
        Generate a .docx document from the letter text.

        Args:
            letter_text: The final letter text
            client_name: Client name for the document
            output_path: Where to save the .docx file
            report_type: Type of report ("monthly_letter" or "workflow_report")

        Returns:
            Dict with output path and metadata
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        if report_type == "monthly_letter":
            self._create_monthly_letter(letter_text, client_name, output_path)
        elif report_type == "workflow_report":
            self._create_workflow_report(letter_text, output_path)

        file_size = os.path.getsize(output_path)
        return {
            "agent": self.name,
            "output_path": output_path,
            "file_size_bytes": file_size,
            "report_type": report_type,
            "success": True,
        }

    def _create_monthly_letter(self, letter_text: str, client_name: str, output_path: str):
        """Create a professionally formatted monthly investment letter (max 2 pages)."""
        doc = Document()

        # --- Set default style to tight spacing ---
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.15

        # Page setup — slightly tighter margins to fit content
        section = doc.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        # --- Compact header: XP branding in one paragraph ---
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_para.paragraph_format.space_after = Pt(0)
        run = header_para.add_run("XP Investimentos")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.name = "Calibri"
        # Add subtitle on same paragraph via line break
        run = header_para.add_run("\n")
        run.font.size = Pt(6)
        run = header_para.add_run("Assessoria de Investimentos")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x7B, 0x8F)
        run.font.name = "Calibri"

        # Separator line
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(2)
        sep.paragraph_format.space_after = Pt(6)
        run = sep.add_run("_" * 90)
        run.font.size = Pt(5)
        run.font.color.rgb = RGBColor(0xED, 0xB9, 0x2E)
        run.font.name = "Calibri"

        # --- Letter body ---
        # Split into paragraphs, skip empty lines (spacing handles separation)
        paragraphs = letter_text.split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                # Skip blank lines — paragraph spacing handles visual breaks
                continue

            # Detect section headers
            is_section = (
                para_text.endswith(":") and len(para_text) < 60
            ) or any(
                para_text.lower().startswith(h)
                for h in [
                    "cenário macro",
                    "desempenho",
                    "posicionamento",
                    "perspectivas",
                    "recomendações",
                    "alocação",
                    "rentabilidade mensal",
                ]
            )

            # Detect bullet items (stock returns, recommendations)
            is_bullet = para_text.startswith("- ")

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15

            if para_text.startswith("Carta Mensal"):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(para_text)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
            elif para_text.startswith("Prezado"):
                run = p.add_run(para_text)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
            elif is_section:
                run = p.add_run(para_text)
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
            elif is_bullet:
                run = p.add_run(para_text)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.line_spacing = 1.0
            elif para_text.startswith("---"):
                run = p.add_run("_" * 90)
                run.font.size = Pt(5)
                run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
            elif para_text.startswith("Atenciosamente"):
                run = p.add_run(para_text)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
            elif "disclaimer" in para_text.lower() or "não constitui" in para_text.lower() or "rentabilidade passada" in para_text.lower():
                run = p.add_run(para_text)
                run.font.size = Pt(7.5)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.italic = True
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
            else:
                run = p.add_run(para_text)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

            # Set font family for all runs
            for run in p.runs:
                run.font.name = "Calibri"

        doc.save(output_path)

    def _create_workflow_report(self, content: str, output_path: str):
        """Create the workflow report answering the 3 challenge questions."""
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("XP AAI Platform — Workflow Report")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6B, 0x7B, 0x8F)
        run.font.name = "Calibri"

        # Separator
        doc.add_paragraph("")

        # Content
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue

            p = doc.add_paragraph()

            if line.startswith("# "):
                run = p.add_run(line[2:])
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            elif line.startswith("## "):
                run = p.add_run(line[3:])
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                p.space_before = Pt(12)
            elif line.startswith("- "):
                run = p.add_run(line)
                run.font.size = Pt(10.5)
                p.paragraph_format.left_indent = Cm(1)
            else:
                run = p.add_run(line)
                run.font.size = Pt(10.5)

            for run in p.runs:
                run.font.name = "Calibri"
                if not run.font.color.rgb:
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        doc.save(output_path)
